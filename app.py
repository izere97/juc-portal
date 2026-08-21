import base64
from datetime import date, datetime
import io
import os
import pandas as pd
from docx import Document
from sqlalchemy import create_engine, text
import streamlit as st
st.markdown(
    '<meta name="google-site-verification" content="IGderbV-0e_PYIBMeJIrTt3uKUtH4Njbq0T7JmWt_OA" />',
    unsafe_allow_html=True,
)
# --- CONFIGURATION ---
st.set_page_config(
    page_title="JUC M&E Portal - Jesuit Urumuri Centre",
    page_icon="background.jpg",
    layout="wide",
)

# --- TRANSLATIONS ---
def get_translations():
    return {
        "English": {
            "title": "JUC Portal", 
            "weekly": "Weekly Report", 
            "strat": "Strategic Pillar Report", 
            "dept": "Department", 
            "submit": "Submit", 
            "lang": "Language", 
            "pillar": "Strategic Pillar", 
            "dash": "📊 Bubble Dashboard", 
            "admin": "⚙️ Admin",
            "capacity": "Ngororero Program",
            "youth_proj": "Youth Innovation (Kigali)"
        },
        "Français": {
            "title": "Portail JUC", 
            "weekly": "Rapport Hebdomadaire", 
            "strat": "Rapport par Pilier Stratégique", 
            "dept": "Département", 
            "submit": "Soumettre", 
            "lang": "Langue", 
            "pillar": "Pilier Stratégique", 
            "dash": "📊 Tableau de Bord en Bulles", 
            "admin": "⚙️ Admin",
            "capacity": "Programme Ngororero",
            "youth_proj": "Innovation Jeunesse (Kigali)"
        },
        "Kinyarwanda": {
            "title": "Urubuga JUC", 
            "weekly": "Raporo y'icyumweru", 
            "strat": "Raporo y'Inkingi z'Ingamba", 
            "dept": "Ishami", 
            "submit": "Ohereza", 
            "lang": "Ururimi", 
            "pillar": "Inkingi y'Ingamba", 
            "dash": "📊 Imbonerahamwe y'Utugari", 
            "admin": "⚙️ Ubuyobozi",
            "capacity": "Gahunda ya Ngororero",
            "youth_proj": "Ihangahanga ry'Urubyiruko (Kigali)"
        },
        "Dutch": {
            "title": "JUC Portaal", 
            "weekly": "Wekelijks Rapport", 
            "strat": "Strategisch Pijler Rapport", 
            "dept": "Afdeling", 
            "submit": "Indienen", 
            "lang": "Taal", 
            "pillar": "Strategische Pijler", 
            "dash": "📊 Bellen Dashboard", 
            "admin": "⚙️ Beheer",
            "capacity": "Ngororero Programma",
            "youth_proj": "Jongeren Innovatie (Kigali)"
        },
        "Italian": {
            "title": "Portale JUC", 
            "weekly": "Rapporto Settimanale", 
            "strat": "Rapporto Pilastro Strategico", 
            "dept": "Dipartimento", 
            "submit": "Invia", 
            "lang": "Lingua", 
            "pillar": "Pilastro Strategico", 
            "dash": "📊 Dashboard a Bolle", 
            "admin": "⚙️ Admin",
            "capacity": "Programma Ngororero",
            "youth_proj": "Innovazione Giovanile (Kigali)"
        },
        "Spanish": {
            "title": "Portal JUC", 
            "weekly": "Informe Semanal", 
            "strat": "Informe de Pilar Estratégico", 
            "dept": "Departamento", 
            "submit": "Enviar", 
            "lang": "Idioma", 
            "pillar": "Pilar Estratégico", 
            "dash": "📊 Panel de Burbujas", 
            "admin": "⚙️ Panel",
            "capacity": "Programa Ngororero",
            "youth_proj": "Innovación Juvenil (Kigali)"
        }
    }

# --- INITIALIZATION ---
if "lang" not in st.session_state: st.session_state.lang = "English"
if "authenticated" not in st.session_state: st.session_state.authenticated = False

trans = get_translations()
engine = create_engine(st.secrets["DATABASE_URL"]) if "DATABASE_URL" in st.secrets else None

# --- INITIALIZE SESSION STATE: NGORORERO PROGRAM ---
if "beneficiaries" not in st.session_state:
    st.session_state.beneficiaries = pd.DataFrame(
        columns=["Full Name", "Sector (Ngororero)", "Training Module", "Phone Number", "Registration Date"]
    )

if "expenses" not in st.session_state:
    st.session_state.expenses = pd.DataFrame(
        {
            "Budget Line": [
                "Selection of Beneficiaries",
                "Vocational Training",
                "Business Management Formation",
                "SILC Support & Equipment",
                "Program Implementation (Salaries)",
            ],
            "Allocated Budget (RWF)": [750000, 31800000, 6120000, 27400000, 16800000],
            "Actual Spent (RWF)": [0, 0, 0, 0, 0],
        }
    )

if "activities" not in st.session_state:
    st.session_state.activities = pd.DataFrame(
        {
            "Activity Name": [
                "Selection of beneficiaries",
                "Vocational training for income-generating initiatives",
                "Formation in resource and business management",
                "Formation of Savings and Internal Lending Communities (SILCs) and support",
                "Program Implementation (Salaries)",
            ],
            "Planned Timeline": ["Q1 - Q2", "Q2 - Q5", "Q3 - Q6", "Q4 - Q8", "Q1 - Q8"],
            "Status": ["Not Started", "Not Started", "Not Started", "Not Started", "Not Started"],
            "Progress Notes": ["", "", "", "", ""],
        }
    )

# --- INITIALIZE SESSION STATE: YOUTH INNOVATION & SOCIAL ENTREPRENEURSHIP (KIGALI) ---
if "youth_beneficiaries" not in st.session_state:
    st.session_state.youth_beneficiaries = pd.DataFrame(
        columns=["Full Name", "Cohort", "Business Idea Title", "District / Suburb", "Phone Number", "Status"]
    )

if "youth_budget" not in st.session_state:
    st.session_state.youth_budget = pd.DataFrame(
        {
            "Budget Line": [
                "1. Recruitment & Selection (Admin/Transport & Competitions)",
                "2. Trainings (Program Mgr, Asst, Mentors, Seed Funding, Equipment, Graduation)",
                "3. Incubation & Support (Facilitation Fees, Mentorship, Wi-Fi, Guest Speakers)"
            ],
            "Allocated Budget (RWF)": [6100000, 92000000, 47600000],
            "Actual Spent (RWF)": [0, 0, 0]
        }
    )

if "youth_activities" not in st.session_state:
    st.session_state.youth_activities = pd.DataFrame(
        {
            "Activity Module / Phase": [
                "Recruitment: Identification of candidates & Competitions",
                "Training: Self-Discovery Module",
                "Training: Self-Realization, Innovation & Prototyping",
                "Training: Marketing and Promotion",
                "Training: Operations, Financing & Financial Management",
                "Training: Strategic Planning and Sustainability",
                "Incubation: Enrolling into incubator & monthly coaching",
                "M&E: Quarterly project evaluations"
            ],
            "Timeline": ["Months 1-3 per cohort", "Phase 1", "Phase 2", "Phase 3", "Phase 4", "Phase 5", "6 Months Incubation", "Continuous (36 Mos)"],
            "Status": ["Not Started", "Not Started", "Not Started", "Not Started", "Not Started", "Not Started", "Not Started", "Not Started"],
            "Remarks": ["", "", "", "", "", "", "", ""]
        }
    )

# --- AUTOMATIC BACKGROUND IMAGE LOADING ---
default_bg_name = "background.jpg"
if "bg_base64" not in st.session_state or not st.session_state.bg_base64:
    if os.path.exists(default_bg_name):
        with open(default_bg_name, "rb") as image_file:
            st.session_state.bg_base64 = base64.b64encode(image_file.read()).decode()
    else:
        st.session_state.bg_base64 = ""

# --- DATABASE STRUCTURE ---
if engine:
    try:
        with engine.connect() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS juc_reports (
                    id SERIAL PRIMARY KEY,
                    submission_date DATE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    staff_name TEXT,
                    report_type TEXT,
                    category TEXT,
                    sub_category TEXT,
                    completed_activities TEXT,
                    pending_issues TEXT,
                    challenges TEXT
                );
            """))
            conn.execute(text("ALTER TABLE juc_reports ADD COLUMN IF NOT EXISTS completed_activities TEXT;"))
            conn.execute(text("ALTER TABLE juc_reports ADD COLUMN IF NOT EXISTS pending_issues TEXT;"))
            conn.execute(text("ALTER TABLE juc_reports ADD COLUMN IF NOT EXISTS challenges TEXT;"))
            conn.commit()
    except Exception as e:
        st.error(f"Database initialization error: {e}")

# --- UI & SIDEBAR AUTHENTICATION ---
st.sidebar.subheader(trans[st.session_state.lang]["lang"])
st.session_state.lang = st.sidebar.selectbox("", ["English", "Français", "Kinyarwanda", "Dutch", "Italian", "Spanish"])
t = trans[st.session_state.lang]

st.sidebar.markdown("---")
st.sidebar.subheader("🔒 Authentication")

if not st.session_state.authenticated:
    with st.sidebar.form("login_form"):
        password_input = st.text_input("Password", type="password")
        submit_button = st.form_submit_button("Sign In")
        
        if submit_button:
            if password_input == "JUC2026Secure":
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("Incorrect password.")

if not st.session_state.authenticated:
    st.warning("Access Restricted. Please enter the secure password in the sidebar and click **Sign In**.")
    st.stop()

if st.sidebar.button("Sign Out"):
    st.session_state.authenticated = False
    st.rerun()

bg_css = f"url('data:image/jpeg;base64,{st.session_state.bg_base64}')" if st.session_state.bg_base64 else "none"

# ... (ton code précédent)

# Application du nouveau style
st.markdown(f"""
    <style>
    /* 1. Fond fixe (Image floue et assombrie) */
    .stApp {{ background: none !important; }}
    .stApp::before {{
        content: ""; position: fixed; top: 0; left: 0; width: 100%; height: 100%;
        background-image: {bg_css} !important; background-size: cover !important;
        background-position: center !important; filter: blur(8px); z-index: -2;
    }}
    .stApp::after {{
        content: ""; position: fixed; top: 0; left: 0; width: 100%; height: 100%;
        background-color: rgba(0, 0, 0, 0.7); z-index: -1;
    }}

    /* 2. Texte global par défaut (pour le fond sombre) */
    div, p, h1, h2, h3, h4, span, label, li {{
        color: #ffffff !important;
        text-shadow: 1px 1px 2px rgba(0,0,0,1);
    }}

    /* 3. Texte noir foncé dans les espaces blancs - CIBLAGE AGRESSIF */
    .main .block-container input, 
    .main .block-container textarea,
    .main .block-container div[role="combobox"],
    .main .block-container div[data-baseweb="select"],
    .main .block-container div[data-baseweb="base-input"] > div,
    .main .block-container div[data-baseweb="textarea"] > textarea {{
        color: #000000 !important; 
        font-weight: 700 !important; 
        background-color: #ffffff !important;
        -webkit-text-fill-color: #000000 !important; /* Force le remplissage du texte */
    }}
    
    /* Ciblage spécifique du placeholder */
    .main .block-container input::placeholder, 
    .main .block-container textarea::placeholder {{
        color: #333333 !important;
        opacity: 1 !important;
    }}

    /* Labels au-dessus des champs blancs */
    .main .block-container label {{
        color: #000000 !important;
        text-shadow: none !important;
        font-weight: bold !important;
    }}

    /* 4. Sidebar en bleu clair */
    [data-testid="stSidebar"] {{ background-color: #f0f9ff !important; }}
    [data-testid="stSidebar"] div, [data-testid="stSidebar"] p, [data-testid="stSidebar"] span {{
        color: #0369a1 !important;
        text-shadow: none !important;
        font-weight: 600 !important;
    }}
    </style>
""", unsafe_allow_html=True)
# ... (le reste de ton code)

# --- NAVIGATION ---
menu = st.sidebar.radio("Navigation", [t["weekly"], t["strat"], t["dash"], t["admin"], t["capacity"], t["youth_proj"]])
# --- 1. WEEKLY REPORT ---
if menu == t["weekly"]:
    st.header(t["weekly"])
    st.info("In accordance with management memo, please submit your weekly activity report and your projections for next week.")
    
    with st.form("weekly_memo_form"):
        col1, col2 = st.columns(2)
        with col1:
            staff_name = st.text_input("Full Name")
        with col2:
            submission_date = st.date_input("Submission Date", value=date.today())
            
        dept = st.selectbox(t["dept"], ["Administration", "Finance", "Program management", "Project office", "Communication office", "Front Desk", "Monitoring and Evaluation"])
        
        completed_activities = st.text_area("Activities completed for the week ending on Friday")
        pending_issues = st.text_area("Projection of pending issues to be completed or initiated next week")
        challenges = st.text_area("Challenges Encountered")
        
        doc = st.file_uploader("Upload supporting document", type=['pdf', 'jpg', 'png', 'docx'])
        
        if st.form_submit_button(t["submit"]):
            if not staff_name or not completed_activities:
                st.error("Please fill in your name and completed activities.")
            elif engine:
                with engine.connect() as conn:
                    conn.execute(text("""
                        INSERT INTO juc_reports (
                            submission_date, staff_name, report_type, category, sub_category, 
                            completed_activities, pending_issues, challenges
                        ) 
                        VALUES (
                            :sub_date, :n, 'Weekly', :c, :sub, 
                            :comp, :pend, :chal
                        )
                    """), {
                        "sub_date": submission_date,
                        "n": staff_name, 
                        "c": dept, 
                        "sub": "JUC weekly report",
                        "comp": completed_activities,
                        "pend": pending_issues,
                        "chal": challenges
                    })
                    conn.commit()
                st.success("Weekly report submitted successfully!")

# --- 2. STRATEGIC PILLAR REPORT ---
elif menu == t["strat"]:
    st.header(t["strat"])
    st.markdown("Please select the relevant strategic pillar to access its specific objectives and activities.")

    chosen_pillar = st.selectbox(
        "Select Strategic Pillar",
        [
            "Pillar 1: Research, Policy Advocacy and Civic Engagement",
            "Pillar 2: Women & Youth Empowerment through Social Innovation and Entrepreneurship",
            "Pillar 3: Integral Ecology and Community Resilience",
            "Pillar 4: Institutional Capacity Strengthening and Sustainability"
        ]
    )
    st.markdown("---")

    if chosen_pillar.startswith("Pillar 1"):
        with st.form("form_pillar_1"):
            staff_name = st.text_input("Full Name", key="p1_name")
            submission_date = st.date_input("Submission Date", value=date.today(), key="p1_date")
            selected_activity = st.selectbox("Core Activity", ["Obj 1.1 - Publication Basic Needs Basket", "Obj 1.2 - Youth Life-skills", "Obj 1.3 - AHAPPY Program"], key="p1_act")
            quantitative_metrics = st.text_input("Quantitative Metrics", key="p1_qm")
            beneficiaries = st.text_input("Beneficiaries", key="p1_ben")
            completed_activities = st.text_area("Progress Details", key="p1_det")
            pending_issues = st.text_area("Pending / Projections for next week", key="p1_pend")
            challenges = st.text_area("Challenges", key="p1_chal")
            if st.form_submit_button(t["submit"]):
                if engine:
                    with engine.connect() as conn:
                        conn.execute(text("""
                            INSERT INTO juc_reports (
                                submission_date, staff_name, report_type, category, sub_category, 
                                completed_activities, pending_issues, challenges
                            ) VALUES (:sub_date, :n, 'Strategic', :c, :sub, :comp, :pend, :chal)
                        """), {
                            "sub_date": submission_date, "n": staff_name, "c": chosen_pillar, "sub": f"{selected_activity} | Metrics: {quantitative_metrics} | Ben: {beneficiaries}",
                            "comp": completed_activities, "pend": pending_issues, "chal": challenges
                        })
                        conn.commit()
                    st.success("Submitted successfully!")
    elif chosen_pillar.startswith("Pillar 2"):
        with st.form("form_pillar_2"):
            staff_name = st.text_input("Full Name", key="p2_name")
            submission_date = st.date_input("Submission Date", value=date.today(), key="p2_date")
            selected_activity = st.selectbox("Core Activity", ["Obj 2.1 - Social innovation incubation", "Obj 2.1 - Financial literacy"], key="p2_act")
            quantitative_metrics = st.text_input("Quantitative Metrics", key="p2_qm")
            beneficiaries = st.text_input("Beneficiaries", key="p2_ben")
            completed_activities = st.text_area("Progress Details", key="p2_det")
            pending_issues = st.text_area("Pending / Projections for next week", key="p2_pend")
            challenges = st.text_area("Challenges", key="p2_chal")
            if st.form_submit_button(t["submit"]):
                if engine:
                    with engine.connect() as conn:
                        conn.execute(text("""
                            INSERT INTO juc_reports (
                                submission_date, staff_name, report_type, category, sub_category, 
                                completed_activities, pending_issues, challenges
                            ) VALUES (:sub_date, :n, 'Strategic', :c, :sub, :comp, :pend, :chal)
                        """), {
                            "sub_date": submission_date, "n": staff_name, "c": chosen_pillar, "sub": f"{selected_activity} | Metrics: {quantitative_metrics} | Ben: {beneficiaries}",
                            "comp": completed_activities, "pend": pending_issues, "chal": challenges
                        })
                        conn.commit()
                    st.success("Submitted successfully!")
    elif chosen_pillar.startswith("Pillar 3"):
        with st.form("form_pillar_3"):
            staff_name = st.text_input("Full Name", key="p3_name")
            submission_date = st.date_input("Submission Date", value=date.today(), key="p3_date")
            selected_activity = st.selectbox("Core Activity", ["Obj 3.1 - Climate change awareness", "Obj 3.2 - Sustainable agriculture"], key="p3_act")
            quantitative_metrics = st.text_input("Quantitative Metrics", key="p3_qm")
            beneficiaries = st.text_input("Beneficiaries", key="p3_ben")
            completed_activities = st.text_area("Progress Details", key="p3_det")
            pending_issues = st.text_area("Pending / Projections for next week", key="p3_pend")
            challenges = st.text_area("Challenges", key="p3_chal")
            if st.form_submit_button(t["submit"]):
                if engine:
                    with engine.connect() as conn:
                        conn.execute(text("""
                            INSERT INTO juc_reports (
                                submission_date, staff_name, report_type, category, sub_category, 
                                completed_activities, pending_issues, challenges
                            ) VALUES (:sub_date, :n, 'Strategic', :c, :sub, :comp, :pend, :chal)
                        """), {
                            "sub_date": submission_date, "n": staff_name, "c": chosen_pillar, "sub": f"{selected_activity} | Metrics: {quantitative_metrics} | Ben: {beneficiaries}",
                            "comp": completed_activities, "pend": pending_issues, "chal": challenges
                        })
                        conn.commit()
                    st.success("Submitted successfully!")
    elif chosen_pillar.startswith("Pillar 4"):
        with st.form("form_pillar_4"):
            staff_name = st.text_input("Full Name", key="p4_name")
            submission_date = st.date_input("Submission Date", value=date.today(), key="p4_date")
            selected_activity = st.selectbox("Core Activity", ["Obj 4.1 - Staff capacity building", "Obj 4.2 - Corporate partnerships"], key="p4_act")
            quantitative_metrics = st.text_input("Quantitative Metrics", key="p4_qm")
            beneficiaries = st.text_input("Beneficiaries", key="p4_ben")
            completed_activities = st.text_area("Progress Details", key="p4_det")
            pending_issues = st.text_area("Pending / Projections for next week", key="p4_pend")
            challenges = st.text_area("Challenges", key="p4_chal")
            if st.form_submit_button(t["submit"]):
                if engine:
                    with engine.connect() as conn:
                        conn.execute(text("""
                            INSERT INTO juc_reports (
                                submission_date, staff_name, report_type, category, sub_category, 
                                completed_activities, pending_issues, challenges
                            ) VALUES (:sub_date, :n, 'Strategic', :c, :sub, :comp, :pend, :chal)
                        """), {
                            "sub_date": submission_date, "n": staff_name, "c": chosen_pillar, "sub": f"{selected_activity} | Metrics: {quantitative_metrics} | Ben: {beneficiaries}",
                            "comp": completed_activities, "pend": pending_issues, "chal": challenges
                        })
                        conn.commit()
                    st.success("Submitted successfully!")

# --- 3. BUBBLE DASHBOARD & EXPORTS (EXCEL & WORD) ---
elif menu == t["dash"]:
    st.header(t["dash"])
    st.markdown("Overview, global exports, and detailed view by department/pillar.")
    
    if engine:
        try:
            df = pd.read_sql("SELECT * FROM juc_reports", engine)
            
            if not df.empty:
                st.subheader("📥 Global Data Export")
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='JUC_Reports')
                
                st.download_button(
                    label="📊 Download all data as Excel (.xlsx)",
                    data=excel_buffer.getvalue(),
                    file_name=f"JUC_Global_Data_Export_{date.today()}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
                st.markdown("---")

            st.markdown("""
                <div class="bubble-admin">
                    <h3>🏛️ Administration & Management</h3>
                    <p>General supervision and institutional steering</p>
                </div>
            """, unsafe_allow_html=True)
            
            col_left, col_mid, col_right = st.columns(3)
            
            with col_left:
                st.markdown("""
                    <div class="bubble-card">
                        <h4>💼 Finance & Admin</h4>
                        <p>Financial management & Operations</p>
                    </div>
                """, unsafe_allow_html=True)
                fin_count = len(df[df['category'] == 'Finance']) if not df.empty and 'category' in df.columns else 0
                st.metric("Finance Reports", fin_count)
            
            with col_mid:
                st.markdown("""
                    <div class="bubble-card">
                        <h4>📊 Program Management</h4>
                        <p>Program coordination & M&E</p>
                    </div>
                """, unsafe_allow_html=True)
                prog_count = len(df[df['category'].str.contains('Program|Monitoring', case=False, na=False)]) if not df.empty and 'category' in df.columns else 0
                st.metric("Program Reports", prog_count)
            
            with col_right:
                st.markdown("""
                    <div class="bubble-card">
                        <h4>🎯 Strategic Pillars</h4>
                        <p>Pillars 1, 2, 3 & 4</p>
                    </div>
                """, unsafe_allow_html=True)
                strat_count = len(df[df['report_type'] == 'Strategic']) if not df.empty and 'report_type' in df.columns else 0
                st.metric("Pillar Reports", strat_count)

            st.markdown("---")
            
            st.subheader("🔍 Detailed Consultation and Word Reports")
            
            if not df.empty:
                export_mode = st.radio("Display and Word Export Mode:", ["Filter by Department or Pillar", "Global Summary of all departments (Synthesis + full details)"])
                
                if export_mode == "Filter by Department or Pillar":
                    all_categories = df['category'].dropna().unique().tolist()
                    selected_cat_view = st.selectbox("Choose department or pillar:", all_categories)
                    
                    filtered_df = df[df['category'] == selected_cat_view]
                    st.markdown(f"### Reports for: **{selected_cat_view}**")
                    
                    for index, row in filtered_df.iterrows():
                        with st.expander(f"👤 {row.get('staff_name', 'N/A')} — Date: {row.get('submission_date', 'N/A')} ({row.get('sub_category', '')})"):
                            st.markdown(f"**✅ Completed Activities:**\n{row.get('completed_activities', 'N/A')}")
                            st.markdown(f"**⏳ Projections / Pending:**\n{row.get('pending_issues', 'N/A')}")
                            st.markdown(f"**⚠️ Challenges Encountered:**\n{row.get('challenges', 'N/A')}")
                    
                    st.markdown("---")
                    
                    if st.button(f"Generate Word document for {selected_cat_view}"):
                        doc = Document()
                        doc.add_heading(f"JUC Report - {selected_cat_view}", 0)
                        doc.add_paragraph(f"Generation date: {date.today().strftime('%Y-%m-%d')}")
                        doc.add_heading("Activity details by department", level=1)
                        
                        for index, row in filtered_df.iterrows():
                            p = doc.add_paragraph()
                            p.add_run(f"Staff Member: {row.get('staff_name', 'N/A')}").bold = True
                            p.add_run(f"\nDate: {row.get('submission_date', 'N/A')}\n")
                            p.add_run(f"Completed Activities:\n{row.get('completed_activities', 'N/A')}\n")
                            p.add_run(f"Projections:\n{row.get('pending_issues', 'N/A')}\n")
                            p.add_run(f"Challenges:\n{row.get('challenges', 'N/A')}\n\n")
                        
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        st.download_button(
                            label=f"📥 Download Word for {selected_cat_view}",
                            data=buffer,
                            file_name=f"JUC_Report_{selected_cat_view.replace(' ', '_')}_{date.today()}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                else:
                    st.markdown("### 📋 Global Summary Overview (All Departments)")
                    summary_counts = df['category'].value_counts().reset_index()
                    summary_counts.columns = ['Department / Pillar', 'Number of Reports']
                    st.dataframe(summary_counts, use_container_width=True)
                    
                    st.markdown("---")
                    if st.button("Generate Word Document: Consolidated Global Summary"):
                        doc = Document()
                        doc.add_heading("JUC Consolidated Report - Global Summary", 0)
                        doc.add_paragraph(f"Generation date: {date.today().strftime('%Y-%m-%d')}")
                        
                        doc.add_heading("1. Complete activity details by department and pillar", level=1)
                        for cat in df['category'].dropna().unique():
                            doc.add_heading(f"Department / Pillar: {cat}", level=2)
                            cat_rows = df[df['category'] == cat]
                            
                            for index, row in cat_rows.iterrows():
                                p = doc.add_paragraph()
                                p.add_run(f"Staff Member: {row.get('staff_name', 'N/A')}").bold = True
                                p.add_run(f" — Date: {row.get('submission_date', 'N/A')}\n")
                                p.add_run(f"Completed Activities:\n{row.get('completed_activities', 'N/A')}\n")
                                if row.get('pending_issues'):
                                    p.add_run(f"Projections / Pending:\n{row.get('pending_issues', 'N/A')}\n")
                                if row.get('challenges'):
                                    p.add_run(f"Challenges Encountered:\n{row.get('challenges', 'N/A')}\n")
                                p.add_run("\n")
                        
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        st.download_button(
                            label="📥 Download Consolidated Word Report (.docx)",
                            data=buffer,
                            file_name=f"JUC_Global_Summary_{date.today()}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
            else:
                st.info("No records found at the moment.")
                
        except Exception as e:
            st.error(f"Error loading dashboard: {e}")

# --- 4. ADMIN ---
elif menu == t["admin"]:
    st.header(t["admin"])
    if engine:
        try:
            df = pd.read_sql("SELECT id, submission_date, staff_name, report_type, category, completed_activities, pending_issues, challenges FROM juc_reports ORDER BY submission_date DESC", engine)
            # --- 4. ADMIN ---
elif menu == t["admin"]:
    st.header(t["admin"])
    if engine:
        try:
            df = pd.read_sql("SELECT id, submission_date, staff_name, report_type, category, completed_activities, pending_issues, challenges FROM juc_reports ORDER BY submission_date DESC", engine)
           # --- 4. ADMIN ---
elif menu == t["admin"]:
    st.header(t["admin"])
    if engine:
        try:
            df = pd.read_sql("SELECT id, submission_date, staff_name, report_type, category, completed_activities, pending_issues, challenges FROM juc_reports ORDER BY submission_date DESC", engine)
         # --- 4. ADMIN ---
elif menu == t["admin"]:
    st.header(t["admin"])
    if engine:
        try:
            df = pd.read_sql("SELECT id, submission_date, staff_name, report_type, category, completed_activities, pending_issues, challenges FROM juc_reports ORDER BY submission_date DESC", engine)
            if df.empty:
                st.info("The database is currently empty.")
            else:
                df["submission_date"] = pd.to_datetime(df["submission_date"])
                
                st.subheader("Filtrer les rapports par date")
                date_selectionnee = st.date_input("Choisir une date pour filtrer les rapports")
                
                df_filtre = df[df["submission_date"].dt.date == date_selectionnee]
                
                if not df_filtre.empty:
                    st.success(f"Affichage des rapports pour le : {date_selectionnee}")
                    st.dataframe(df_filtre, use_container_width=True)
                else:
                    st.warning(f"Aucun rapport trouvé pour la date : {date_selectionnee}")
                
                if st.button("Afficher tous les rapports"):
                    st.dataframe(df, use_container_width=True)

            st.subheader("Delete a specific record by ID")
            report_id_to_delete = st.number_input("Enter ID to delete", min_value=0, step=1)

            if st.button("Delete Selected Row"):
                with engine.connect() as conn:
                    conn.execute(text("DELETE FROM juc_reports WHERE id = :id"), {"id": report_id_to_delete})
                    conn.commit()
                    st.success(f"Record ID {report_id_to_delete} deleted successfully.")
                    st.rerun()
                    
        except Exception as e:
            st.info("Loading administration tools...")

        st.markdown("---")
        if st.button("🗑️ DELETE ALL (Total Reset)"):
            with engine.connect() as conn:
                conn.execute(text("DELETE FROM juc_reports"))
                conn.commit()
                st.success("All records deleted successfully.")
                st.rerun()

            # ... (la suite de votre code pour supprimer reste inchangée)
                
                st.subheader("Delete a specific record by ID")
                report_id_to_delete = st.number_input("Enter ID to delete", min_value=0, step=1)
                
                if st.button("Delete Selected Row"):
                    with engine.connect() as conn:
                        conn.execute(text("DELETE FROM juc_reports WHERE id = :id"), {"id": report_id_to_delete})
                        conn.commit()
                    st.success(f"Record ID {report_id_to_delete} deleted successfully.")
                    st.rerun()
        except Exception as e:
            st.info("Loading administration tools...")
            
        st.markdown("---")
        if st.button("🗑️ DELETE ALL (Total Reset)"):
            with engine.connect() as conn:
                conn.execute(text("DELETE FROM juc_reports"))
                conn.commit()
            st.warning("Database completely cleared.")
            st.rerun()

# --- 5. CAPACITY BUILDING PROGRAM (NGORORERO DISTRICT) ---
elif menu == t["capacity"]:
    st.title("Capacity Building and Social Empowerment Program")
    st.subheader("Ngororero District - Monitoring & Evaluation Platform")
    st.markdown("---")

    cap_nav = st.sidebar.radio(
        "Ngororero Sections",
        [
            "Beneficiary Management",
            "Financial & Budget Tracker",
            "Activity Monitoring",
            "Project Overview",
        ],
    )

    if cap_nav == "Beneficiary Management":
        st.header("Beneficiary Registration & Tracking")
        st.markdown("Register new vulnerable women beneficiaries in Ngororero District.")

        with st.form("beneficiary_form"):
            col1, col2 = st.columns(2)
            with col1:
                full_name = st.text_input("Full Name")
                sector = st.text_input("Ngororero Sector")
            with col2:
                training_module = st.selectbox(
                    "Vocational Training Path",
                    ["Hairdressing", "Culinary Arts", "Tailoring", "Logistics"],
                )
                phone = st.text_input("Phone Number")

            submit_btn = st.form_submit_button("Register Beneficiary")

            if submit_btn:
                if full_name and sector:
                    new_row = pd.DataFrame(
                        {
                            "Full Name": [full_name],
                            "Sector (Ngororero)": [sector],
                            "Training Module": [training_module],
                            "Phone Number": [phone],
                            "Registration Date": [pd.Timestamp.today().strftime("%Y-%m-%d")],
                        }
                    )
                    st.session_state.beneficiaries = pd.concat(
                        [st.session_state.beneficiaries, new_row], ignore_index=True
                    )
                    st.success(f"Successfully registered {full_name}!")
                else:
                    st.error("Please fill in at least the Full Name and Sector.")

        st.markdown("---")
        st.subheader("Registered Beneficiaries List")
        if not st.session_state.beneficiaries.empty:
            st.dataframe(st.session_state.beneficiaries, use_container_width=True)
            total_ben = len(st.session_state.beneficiaries)
            st.metric(label="Total Registered Beneficiaries", value=total_ben)
        else:
            st.info("No beneficiaries registered yet.")

    elif cap_nav == "Financial & Budget Tracker":
        st.header("Financial Capacity & Budget Utilization")
        st.markdown("Monitor allocated budgets versus actual expenditures across project lines.")

        st.subheader("Current Budget Overview (RWF)")
        df_expenses = st.session_state.expenses
        df_expenses["Remaining Budget (RWF)"] = (
            df_expenses["Allocated Budget (RWF)"]
            - df_expenses["Actual Spent (RWF)"]
        )
        df_expenses["Utilization (%)"] = (
            df_expenses["Actual Spent (RWF)"]
            / df_expenses["Allocated Budget (RWF)"]
            * 100
        ).round(2)

        st.dataframe(df_expenses, use_container_width=True)

        st.markdown("### Log an Expense")
        with st.form("expense_form"):
            selected_line = st.selectbox(
                "Select Budget Line", df_expenses["Budget Line"]
            )
            expense_amount = st.number_input(
                "Amount Spent (RWF)", min_value=0.0, step=1000.0
            )
            log_btn = st.form_submit_button("Record Expense")

            if log_btn:
                idx = st.session_state.expenses[
                    st.session_state.expenses["Budget Line"] == selected_line
                ].index[0]
                st.session_state.expenses.loc[idx, "Actual Spent (RWF)"] += expense_amount
                st.rerun()

        total_allocated = df_expenses["Allocated Budget (RWF)"].sum()
        total_spent = df_expenses["Actual Spent (RWF)"].sum()

        col1, col2, col3 = st.columns(3)
        col1.metric("Total Budget", f"{total_allocated:,.0f} RWF")
        col2.metric("Total Spent", f"{total_spent:,.0f} RWF")
        col3.metric(
            "Overall Utilization",
            f"{(total_spent / total_allocated * 100):.2f}%"
            if total_allocated > 0
            else "0%",
        )

    elif cap_nav == "Activity Monitoring":
        st.header("Project Activity Monitoring")
        st.markdown("Track execution status and progress notes for planned activities.")

        st.subheader("Master Activity Schedule & Status")
        st.dataframe(st.session_state.activities, use_container_width=True)

        st.markdown("### Update Activity Status")
        with st.form("activity_form"):
            selected_activity = st.selectbox(
                "Select Activity", st.session_state.activities["Activity Name"]
            )
            new_status = st.selectbox(
                "Update Status", ["Not Started", "In Progress", "Completed"]
            )
            new_notes = st.text_area("Progress Notes / Remarks")
            update_btn = st.form_submit_button("Save Activity Update")

            if update_btn:
                idx = st.session_state.activities[
                    st.session_state.activities["Activity Name"] == selected_activity
                ].index[0]
                st.session_state.activities.loc[idx, "Status"] = new_status
                st.session_state.activities.loc[idx, "Progress Notes"] = new_notes
                st.success("Activity status updated successfully!")
                st.rerun()

    elif cap_nav == "Project Overview":
        st.header("Project Summary & Metadata")
        st.markdown(
            """
            * **Project Name:** Capacity Building and Social Empowerment Program for Vulnerable Women
            * **Location:** Ngororero District, Rwanda
            * **Start Date:** April 2026
            * **Implementing Body:** Jesuit Urumuri Centre (JUG)
            * **Total Approved Budget:** 82,870,000 RWF (53,433 €)
            """
        )
        st.info("Use the sidebar sub-options to switch between beneficiary management, financial tracking, and activity monitoring.")

# --- 6. YOUTH INNOVATION & SOCIAL ENTREPRENEURSHIP (KIGALI) ---
elif menu == t["youth_proj"]:
    st.title("Youth Innovation and Social Entrepreneurship Project")
    st.subheader("City of Kigali Suburbs — 36-Month 10-Cohort Program (CEI Partnership)")
    st.markdown("---")

    youth_nav = st.sidebar.radio(
        "Kigali Youth Project Sections",
        [
            "Beneficiaries & Cohorts",
            "Financial & General Budget",
            "Action Learning & Activities",
            "Project Proposal Overview"
        ]
    )

    if youth_nav == "Beneficiaries & Cohorts":
        st.header("Youth Beneficiary Tracking & Cohort Management")
        st.markdown("Register university/college graduate youth from Kigali suburbs for upcoming training cohorts[cite: 1].")

        with st.form("youth_beneficiary_form"):
            col1, col2 = st.columns(2)
            with col1:
                y_name = st.text_input("Full Name")
                y_cohort = st.selectbox("Cohort Intake", [f"Cohort {i}" for i in range(1, 11)])
                y_district = st.text_input("District / Suburb (e.g., Gasabo, Kicukiro)")
            with col2:
                y_idea = st.text_input("Business Idea Title")
                y_phone = st.text_input("Phone Number")
                y_status = st.selectbox("Selection Status", ["Candidate", "Selected for Training", "Incubation Stage", "Graduated"])

            y_submit = st.form_submit_button("Register Youth Innovator")

            if y_submit:
                if y_name and y_idea:
                    new_youth_row = pd.DataFrame(
                        {
                            "Full Name": [y_name],
                            "Cohort": [y_cohort],
                            "Business Idea Title": [y_idea],
                            "District / Suburb": [y_district],
                            "Phone Number": [y_phone],
                            "Status": [y_status]
                        }
                    )
                    st.session_state.youth_beneficiaries = pd.concat(
                        [st.session_state.youth_beneficiaries, new_youth_row], ignore_index=True
                    )
                    st.success(f"Successfully added youth innovator: {y_name}!")
                else:
                    st.error("Please provide at least the Full Name and Business Idea Title.")

        st.markdown("---")
        st.subheader("Registered Youth Innovators Registry")
        if not st.session_state.youth_beneficiaries.empty:
            st.dataframe(st.session_state.youth_beneficiaries, use_container_width=True)
            st.metric(label="Total Registered Youth Trainees", value=len(st.session_state.youth_beneficiaries))
        else:
            st.info("No youth beneficiaries registered yet. Target is 400 youth across 10 cohorts (40 per cohort)[cite: 1].")

    elif youth_nav == "Financial & General Budget":
        st.header("Financial Breakdown & Budget Tracking")
        st.markdown("Tracking general budget allocations (CEI Contribution & Local Contribution) totaling **140,300,000 RWF (€118,898)**[cite: 1].")

        st.subheader("Budget Allocation vs. Actual Spending (RWF)")
        df_yb = st.session_state.youth_budget
        df_yb["Remaining Budget (RWF)"] = df_yb["Allocated Budget (RWF)"] - df_yb["Actual Spent (RWF)"]
        df_yb["Utilization (%)"] = (df_yb["Actual Spent (RWF)"] / df_yb["Allocated Budget (RWF)"] * 100).round(2)

        st.dataframe(df_yb, use_container_width=True)

        st.markdown("### Log Expense for Youth Project")
        with st.form("youth_expense_form"):
            selected_line_y = st.selectbox("Select Budget Line", df_yb["Budget Line"])
            y_amt = st.number_input("Amount Spent (RWF)", min_value=0.0, step=5000.0)
            y_log_btn = st.form_submit_button("Record Project Expense")

            if y_log_btn:
                idx = st.session_state.youth_budget[st.session_state.youth_budget["Budget Line"] == selected_line_y].index[0]
                st.session_state.youth_budget.loc[idx, "Actual Spent (RWF)"] += y_amt
                st.success("Expense updated successfully!")
                st.rerun()

        tot_alloc = df_yb["Allocated Budget (RWF)"].sum()
        tot_spent = df_yb["Actual Spent (RWF)"].sum()

        col1, col2, col3 = st.columns(3)
        col1.metric("Total Project Budget", f"{tot_alloc:,.0f} RWF")
        col2.metric("Total Spent", f"{tot_spent:,.0f} RWF")
        col3.metric("Utilization Rate", f"{(tot_spent / tot_alloc * 100):.2f}%" if tot_alloc > 0 else "0%")

    elif youth_nav == "Action Learning & Activities":
        st.header("Action Learning & Incubation Schedule")
        st.markdown("Track the progress of recruitment, the 5 training modules, and 6-month incubations[cite: 1].")

        st.dataframe(st.session_state.youth_activities, use_container_width=True)

        st.markdown("### Update Module / Activity Status")
        with st.form("youth_act_form"):
            sel_act = st.selectbox("Select Phase / Module", st.session_state.youth_activities["Activity Module / Phase"])
            new_st = st.selectbox("Status", ["Not Started", "In Progress", "Completed"], key="y_st")
            new_rk = st.text_area("Remarks / Implementation Notes", key="y_rk")
            upd_btn = st.form_submit_button("Update Status")

            if upd_btn:
                idx = st.session_state.youth_activities[st.session_state.youth_activities["Activity Module / Phase"] == sel_act].index[0]
                st.session_state.youth_activities.loc[idx, "Status"] = new_st
                st.session_state.youth_activities.loc[idx, "Remarks"] = new_rk
                st.success("Activity progress updated successfully!")
                st.rerun()

    elif youth_nav == "Project Overview":
        st.header("Project Proposal Summary (CEI Partnership)")
        st.markdown(
            """
            * **Project Title:** Youth Innovation and Social Entrepreneurship[cite: 1]
            * **Implementing Body:** Jesuit Urumuri Centre (JUC), Kigali, Gasabo District[cite: 1]
            * **Partner:** Conferenza Episcopale Italiana (CEI)[cite: 1]
            * **Target Beneficiaries:** 400 unemployed university and college graduates in Kigali suburbs (10 cohorts of 40 youth each)[cite: 1].
            * **Lifespan:** 36 Months (3 Years)[cite: 1]
            * **Total Budget:** 140,300,000 RWF (€118,898)[cite: 1]
            * **Core Pillars:** Recruitment & Selection, Action Learning (5 Training Modules), and 6-Month Project Incubation[cite: 1].
            """
        )
        st.info("Use the sidebar sub-options to manage youth innovator cohorts, expenses, and action learning modules.")
